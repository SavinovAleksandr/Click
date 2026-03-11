# -*- coding: utf-8 -*-
"""
Click_word - Вставка изображений (PNG/SVG) в документ Word.
Поддерживает python-docx и COM Word для SVG.
"""
import logging
from os import path as os_path
from pathlib import Path

import win32com.client as win32
from docx import Document

from config import (
    STYLE_HEADING_APP, STYLE_ILLUSTRATION, STYLE_ILLUSTRATION_APP,
    STYLE_ILLUSTRATION_LAYOUT, OUTPUT_FILENAME
)

logger = logging.getLogger(__name__)
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION


def orientation_format(section, position_value, format_value):
    """Настройка ориентации и размера страницы."""
    key = (position_value == 'Альбомная', format_value == 'A3')
    page_settings = {
        (False, False): {'orientation': WD_ORIENTATION.PORTRAIT, 'width': Cm(21), 'height': Cm(29.7)},
        (True, False): {'orientation': WD_ORIENTATION.LANDSCAPE, 'width': Cm(29.7), 'height': Cm(21)},
        (False, True): {'orientation': WD_ORIENTATION.PORTRAIT, 'width': Cm(29.7), 'height': Cm(42)},
        (True, True): {'orientation': WD_ORIENTATION.LANDSCAPE, 'width': Cm(42), 'height': Cm(29.7)},
    }
    settings = page_settings.get(key, page_settings[(False, False)])
    section.orientation = settings['orientation']
    section.page_width = settings['width']
    section.page_height = settings['height']


def process_directory(rg2, path_wrd, format_value, position_value, text_value):
    """
    Помещает созданные PNG-файлы в Word, применяет стиль и подписывает.

    Args:
        rg2: список путей к rg2 файлам (для определения имён изображений)
        path_wrd: путь к шаблону Word
        format_value: A4 или A3
        position_value: Книжная или Альбомная
        text_value: использовать стили приложения
    """
    doc = Document(path_wrd)
    section = doc.sections[0]
    orientation_format(section, position_value, format_value)

    width_map = {(False, False): Cm(14), (True, False): Cm(14), (False, True): Cm(14), (True, True): Cm(14)}
    key = (position_value == 'Альбомная', format_value == 'A3')
    width_cm = width_map.get(key, Cm(14))

    filenames = []
    for filename in rg2:
        filenames.append(os_path.join(os_path.dirname(filename), os_path.splitext(filename)[0] + '.png'))

    for filename in filenames:
        if not os_path.exists(filename):
            logger.warning('Файл не найден, пропуск: %s', filename)
            continue
        p = doc.add_paragraph()
        p.style = STYLE_HEADING_APP
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filename, width=width_cm)
        p = doc.add_paragraph()
        if text_value:
            p.style = STYLE_ILLUSTRATION_APP
        else:
            p.style = STYLE_ILLUSTRATION
        tire = ' / '
        p.add_run().add_text(os_path.basename(filename).replace('/', tire))
        doc.add_page_break()

    word_output = os_path.join(os_path.dirname(rg2[0]), OUTPUT_FILENAME)
    doc.save(word_output)


def process_directory_com(rg2, path_wrd, format_value, position_value, text_value):
    """Вставка SVG через COM Word (требует установленный Microsoft Word)."""
    word_app = win32.Dispatch('Word.Application')
    word_app.Visible = False
    word_app.ScreenUpdating = False
    doc = word_app.Documents.Open(path_wrd)

    width_map = {(False, False): 28.35 * 14, (True, False): 28.35 * 14, (False, True): 28.35 * 14, (True, True): 28.35 * 14}

    for rg2_path in rg2:
        directory = os_path.dirname(rg2_path)
        base_name = Path(rg2_path).stem
        filenames = [str(Path(directory) / (base_name + '.svg'))]

        for filename in filenames:
            if not os_path.exists(filename):
                logger.warning('Файл не найден: %s', filename)
                continue
            selection = word_app.Selection
            selection.EndKey(Unit=6)
            inline_shape = selection.InlineShapes.AddPicture(
                filename, LinkToFile=False, SaveWithDocument=True, Width=width_map.get((0, 0), 400)
            )
            selection.Style = STYLE_ILLUSTRATION_LAYOUT
            selection.InsertParagraphAfter()
            selection.MoveDown(Unit=1, Count=1)
            file_name_without_ext = Path(filename).stem
            caption_text = file_name_without_ext.replace('/', ' - ')
            selection.Text = caption_text
            selection.Style = STYLE_ILLUSTRATION if not text_value else STYLE_ILLUSTRATION_APP
            selection.InsertBreak(7)

    word_output = os_path.join(os_path.dirname(rg2[0]), OUTPUT_FILENAME)
    try:
        doc.SaveAs(str(word_output))
    except Exception as e:
        logger.error('Ошибка при работе с документом: %s', e)
    finally:
        doc.Close()
        word_app.Quit()


if __name__ == '__main__':
    process_directory_com(
        [r'C:\Users\Gurev-AA\Desktop\тест\Вот_ГЭС2\.28_1_Р_11_ФПТ_0.dxf'],
        r'C:\Users\Gurev-AA\Desktop\тест\Шаблон НТЦ ЕЭС.docx',
        'A4', 'Книжная', False
    )
